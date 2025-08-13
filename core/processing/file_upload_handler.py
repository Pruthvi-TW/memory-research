"""
File Upload Handler for processing various file types.
Supports PDF, DOCX, TXT, and Markdown files with text extraction.
"""

import io
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import mimetypes

# PDF processing
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    try:
        import pdfplumber
        PDF_AVAILABLE = True
        USE_PDFPLUMBER = True
    except ImportError:
        PDF_AVAILABLE = False
        USE_PDFPLUMBER = False

# DOCX processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


class FileUploadHandler:
    """
    Handler for processing uploaded files and extracting text content.
    """
    
    SUPPORTED_TYPES = {
        'text/plain': 'txt',
        'text/markdown': 'md',
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx'
    }
    
    def __init__(self):
        self.max_file_size = 10 * 1024 * 1024  # 10MB
        self.max_files_per_batch = 10
        
        # Initialize processors
        self.processors = {
            'txt': self._process_text_file,
            'md': self._process_markdown_file,
            'pdf': self._process_pdf_file if PDF_AVAILABLE else None,
            'docx': self._process_docx_file if DOCX_AVAILABLE else None
        }
        
        logger.info("📁 FileUploadHandler initialized")
        if not PDF_AVAILABLE:
            logger.warning("⚠️ PDF processing not available - install PyPDF2 or pdfplumber")
        if not DOCX_AVAILABLE:
            logger.warning("⚠️ DOCX processing not available - install python-docx")
    
    async def process_uploaded_files(self, files: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Process multiple uploaded files and extract text content.
        
        Args:
            files: List of file dictionaries with 'content', 'filename', 'content_type'
            
        Returns:
            List of processed documents with extracted text
        """
        processed_documents = []
        
        if len(files) > self.max_files_per_batch:
            raise ValueError(f"Too many files. Maximum {self.max_files_per_batch} files allowed per batch")
        
        for file_data in files:
            try:
                processed_doc = await self._process_single_file(file_data)
                if processed_doc:
                    processed_documents.append(processed_doc)
                    logger.info(f"✅ Processed file: {file_data.get('filename', 'unknown')}")
                else:
                    logger.warning(f"⚠️ Failed to process file: {file_data.get('filename', 'unknown')}")
                    
            except Exception as e:
                logger.error(f"❌ Error processing file {file_data.get('filename', 'unknown')}: {e}")
                continue
        
        logger.info(f"📊 Processed {len(processed_documents)} out of {len(files)} files")
        return processed_documents
    
    async def _process_single_file(self, file_data: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        """Process a single uploaded file."""
        filename = file_data.get('filename', 'unknown')
        content_type = file_data.get('content_type', '')
        file_content_raw = file_data.get('content', '')
        
        # Handle base64 encoded content from frontend
        if isinstance(file_content_raw, str):
            import base64
            try:
                file_content = base64.b64decode(file_content_raw)
            except Exception as e:
                raise ValueError(f"Failed to decode base64 content: {e}")
        else:
            file_content = file_content_raw
        
        # Validate file
        validation_result = await self.validate_file(file_data)
        if not validation_result['valid']:
            raise ValueError(f"File validation failed: {', '.join(validation_result['errors'])}")
        
        # Determine file type
        file_extension = self.SUPPORTED_TYPES.get(content_type)
        if not file_extension:
            # Try to determine from filename
            file_extension = Path(filename).suffix.lower().lstrip('.')
            if file_extension not in self.processors:
                raise ValueError(f"Unsupported file type: {content_type}")
        
        # Extract text content
        processor = self.processors.get(file_extension)
        if not processor:
            raise ValueError(f"No processor available for {file_extension} files")
        
        extracted_text = await processor(file_content, filename)
        
        if not extracted_text or len(extracted_text.strip()) < 10:
            raise ValueError("Extracted text is too short or empty")
        
        # Create processed document
        processed_doc = {
            'content': extracted_text,
            'source': filename,
            'source_type': 'upload',
            'content_type': content_type,
            'file_extension': file_extension,
            'metadata': {
                'original_filename': filename,
                'file_size': len(file_content),
                'content_type': content_type,
                'extraction_method': file_extension,
                'character_count': len(extracted_text),
                'word_count': len(extracted_text.split())
            }
        }
        
        return processed_doc
    
    async def validate_file(self, file_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Validate uploaded file for security and format compliance.
        
        Args:
            file_data: File data dictionary
            
        Returns:
            Validation result with 'valid', 'errors', 'warnings'
        """
        result = {'valid': True, 'errors': [], 'warnings': []}
        
        filename = file_data.get('filename', '')
        content_type = file_data.get('content_type', '')
        file_content_raw = file_data.get('content', '')
        
        # Handle base64 encoded content from frontend
        if isinstance(file_content_raw, str):
            import base64
            try:
                file_content = base64.b64decode(file_content_raw)
            except Exception:
                file_content = file_content_raw.encode('utf-8')  # Fallback for text content
        else:
            file_content = file_content_raw
        
        # Check file size
        file_size = len(file_content)
        if file_size == 0:
            result['valid'] = False
            result['errors'].append("File is empty")
        elif file_size > self.max_file_size:
            result['valid'] = False
            result['errors'].append(f"File size ({file_size} bytes) exceeds maximum ({self.max_file_size} bytes)")
        
        # Check filename
        if not filename:
            result['valid'] = False
            result['errors'].append("Filename is required")
        elif len(filename) > 255:
            result['valid'] = False
            result['errors'].append("Filename is too long")
        
        # Check for suspicious file extensions
        suspicious_extensions = ['.exe', '.bat', '.cmd', '.scr', '.vbs', '.js', '.jar']
        file_ext = Path(filename).suffix.lower()
        if file_ext in suspicious_extensions:
            result['valid'] = False
            result['errors'].append(f"Potentially unsafe file extension: {file_ext}")
        
        # Validate content type
        if content_type not in self.SUPPORTED_TYPES:
            # Try to guess from filename
            guessed_type, _ = mimetypes.guess_type(filename)
            if guessed_type and guessed_type in self.SUPPORTED_TYPES:
                result['warnings'].append(f"Content type mismatch. Using guessed type: {guessed_type}")
            else:
                result['valid'] = False
                result['errors'].append(f"Unsupported content type: {content_type}")
        
        # Basic content validation
        if file_content:
            # Check for binary content in text files
            if content_type.startswith('text/'):
                try:
                    file_content.decode('utf-8')
                except UnicodeDecodeError:
                    result['warnings'].append("File may contain binary data")
        
        return result
    
    async def _process_text_file(self, file_content: bytes, filename: str) -> str:
        """Process plain text files."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    # Clean up the text
                    text = text.replace('\r\n', '\n').replace('\r', '\n')
                    return text.strip()
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Could not decode text file with any supported encoding")
            
        except Exception as e:
            logger.error(f"❌ Text file processing error: {e}")
            raise
    
    async def _process_markdown_file(self, file_content: bytes, filename: str) -> str:
        """Process Markdown files."""
        try:
            # Markdown files are essentially text files
            text = await self._process_text_file(file_content, filename)
            
            # Basic markdown cleanup - remove some markdown syntax for better text extraction
            import re
            
            # Remove markdown headers but keep the text
            text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
            
            # Remove markdown links but keep the text
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            
            # Remove markdown emphasis but keep the text
            text = re.sub(r'\*\*([^\*]+)\*\*', r'\1', text)
            text = re.sub(r'\*([^\*]+)\*', r'\1', text)
            
            # Remove code blocks
            text = re.sub(r'```[^`]*```', '', text, flags=re.DOTALL)
            text = re.sub(r'`([^`]+)`', r'\1', text)
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"❌ Markdown file processing error: {e}")
            raise
    
    async def _process_pdf_file(self, file_content: bytes, filename: str) -> str:
        """Process PDF files."""
        if not PDF_AVAILABLE:
            raise ValueError("PDF processing not available. Install PyPDF2 or pdfplumber")
        
        try:
            text_content = []
            
            if 'USE_PDFPLUMBER' in globals() and USE_PDFPLUMBER:
                # Use pdfplumber for better text extraction
                import pdfplumber
                
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content.append(page_text)
            else:
                # Use PyPDF2
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
            
            if not text_content:
                raise ValueError("No text content extracted from PDF")
            
            # Join all pages and clean up
            full_text = '\n\n'.join(text_content)
            
            # Clean up common PDF artifacts
            import re
            full_text = re.sub(r'\s+', ' ', full_text)  # Normalize whitespace
            full_text = re.sub(r'\n\s*\n', '\n\n', full_text)  # Clean up line breaks
            
            return full_text.strip()
            
        except Exception as e:
            logger.error(f"❌ PDF processing error: {e}")
            raise
    
    async def _process_docx_file(self, file_content: bytes, filename: str) -> str:
        """Process DOCX files."""
        if not DOCX_AVAILABLE:
            raise ValueError("DOCX processing not available. Install python-docx")
        
        try:
            # Load document from bytes
            doc = Document(io.BytesIO(file_content))
            
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            if not text_content:
                raise ValueError("No text content extracted from DOCX")
            
            # Join all content
            full_text = '\n\n'.join(text_content)
            
            return full_text.strip()
            
        except Exception as e:
            logger.error(f"❌ DOCX processing error: {e}")
            raise
    
    def get_supported_types(self) -> Dict[str, str]:
        """Get supported file types and their descriptions."""
        return {
            'text/plain': 'Plain text files (.txt)',
            'text/markdown': 'Markdown files (.md)',
            'application/pdf': 'PDF documents (.pdf)' if PDF_AVAILABLE else 'PDF documents (.pdf) - Not available',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'Word documents (.docx)' if DOCX_AVAILABLE else 'Word documents (.docx) - Not available'
        }
    
    def get_file_info(self, file_data: Dict[str, Any]) -> Dict[str, Any]:
        """Get information about a file without processing it."""
        filename = file_data.get('filename', 'unknown')
        content_type = file_data.get('content_type', '')
        file_content = file_data.get('content', b'')
        
        return {
            'filename': filename,
            'content_type': content_type,
            'file_size': len(file_content),
            'file_extension': Path(filename).suffix.lower(),
            'supported': content_type in self.SUPPORTED_TYPES,
            'processor_available': self.processors.get(self.SUPPORTED_TYPES.get(content_type)) is not None
        }